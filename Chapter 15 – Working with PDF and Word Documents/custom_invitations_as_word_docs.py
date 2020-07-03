import docx

doc = docx.Document('invitations.docx')
file = open('guests.txt')
guests = file.readlines()


for guest in guests:
    name = guest.replace('\n','')

    para1 = doc.add_paragraph('It would be a pleasure to have the company of')
    para2 = doc.add_paragraph(name)
    para3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    para4 = doc.add_paragraph('Aprill 1st')
    para5 = doc.add_paragraph('at 7o\'clock')
    para5.runs[0].add_break(docx.text.run.WD_BREAK.PAGE)

    para1.style = para3.style = para5.style = 'text_style'
    para2.style = 'name_style'
    para4.style = 'date_style'


doc.save('invitations.docx')